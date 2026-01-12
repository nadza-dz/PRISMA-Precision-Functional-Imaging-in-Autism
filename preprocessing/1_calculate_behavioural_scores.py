#!/usr/bin/env python3
"""
Behavioural Data Processing Script for ASD Precision Study

This script processes behavioural questionnaire data including:
- AQ-50 (Autism Quotient) questionnaire scoring
- SDQ (Strengths and Difficulties Questionnaire) scoring  
- Demographics data processing
- MRI acquisition date extraction
- IQ (WAIS IV Matrix Reasoning and Vocabulary tests) scoring
- ARSQ (Amsterdam Resting State Questionnaire) scoring
- Data merging and final output generation

Author: Joe Bathelt
Date: October 2025
Version: 1.0

ND : added IQ and ARSQ score calculation
Dec 2025
"""

# %%
# Import required libraries
import numpy as np
import os
import pandas as pd
from pathlib import Path
import re
from docx import Document
import csv



# %%
# Setup project paths
project_folder = Path(__file__).resolve().parents[2]
data_folder = project_folder / 'data' / 'raw_data' / 'behavioural'

print(f"Project folder: {project_folder}")
print(f"Data folder: {data_folder}")
# %%
def process_aq_questionnaire(data_folder, project_folder):
    """
    Process AQ-50 (Autism Quotient) questionnaire data.
    
    The AQ-50 is a self-report questionnaire measuring autistic traits in adults.
    It consists of 50 items across 5 domains: social skills, attention switching,
    attention to detail, communication, and imagination.
    
    Parameters:
    -----------
    data_folder : Path
        Path to the behavioural data folder
    project_folder : Path  
        Path to the project root folder
        
    Returns:
    --------
    pd.DataFrame
        Processed AQ scores with subscales and total score
    """
    print("Processing AQ-50 questionnaire data...")
    
    # Load and concatenate all AQ files
    aq_files = [f for f in data_folder.glob('**/AQ*.csv')]
    print(f"Found {len(aq_files)} AQ files")
    
    aq_df = pd.concat([pd.read_csv(f, skiprows=(1,2), index_col='ParticipantID', na_values='NaN') for f in aq_files])
    aq_df = aq_df[[col for col in aq_df.columns if 'AQ' in col]].dropna().copy()
    aq_df.columns = ['AQ' + str(i+1) for i in range(len(sorted(aq_df.columns)))]
    aq_df = aq_df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
    
    # Define forward and reverse scored items
    forward_items = [1, 3, 8, 10, 11, 14, 15, 17, 24, 25, 27, 28, 29, 30, 31, 32, 34, 36, 37, 38, 40, 44, 47, 48, 49, 50]
    reverse_items = [2, 4, 5, 6, 7, 9, 12, 13, 16, 18, 19, 20, 21, 22, 23, 26, 33, 35, 39, 41, 42, 43, 45, 46]

    # Score forward items (higher agreement = higher score)
    for item in forward_items:
        aq_df.loc[:, 'AQ' + str(item)] = aq_df['AQ' + str(item)].replace({
            'geheel mee eens':1,
            'enigszins mee eens':2,
            'enigszins mee oneens':3,
            'geheel mee oneens':4
            })
    
    # Score reverse items (higher agreement = lower score)
    for item in reverse_items:
        aq_df.loc[:, 'AQ' + str(item)] = aq_df['AQ' + str(item)].replace({
            'geheel mee eens':4,
            'enigszins mee eens':3,
            'enigszins mee oneens':2,
            'geheel mee oneens':1
            })
    
    # Define subscale item groupings
    social = [1, 11, 13, 15, 22, 36, 44, 45, 47, 48]
    attention_switching = [2, 4, 10, 16, 25, 32, 34, 37, 43, 46]
    attention_to_detail = [5, 6, 9, 12, 19, 23, 28, 29, 30, 49]
    communication = [7, 17, 18, 26, 27, 31, 33, 35, 38, 39]
    imagination = [3, 8, 14, 20, 21, 24, 40, 41, 42, 50]

    # Calculate subscale and total scores
    aq_df['AQ_social'] = aq_df[[f'AQ{i}' for i in social]].sum(axis=1)
    aq_df['AQ_attention_switching'] = aq_df[[f'AQ{i}' for i in attention_switching]].sum(axis=1)
    aq_df['AQ_attention_to_detail'] = aq_df[[f'AQ{i}' for i in attention_to_detail]].sum(axis=1)
    aq_df['AQ_communication'] = aq_df[[f'AQ{i}' for i in communication]].sum(axis=1)
    aq_df['AQ_imagination'] = aq_df[[f'AQ{i}' for i in imagination]].sum(axis=1)
    aq_df['AQ_total'] = aq_df[[f'AQ{i}' for i in range(1, 51)]].sum(axis=1)
    aq_df['AQ_above_diagnostic_threshold'] = np.where(aq_df['AQ_total'] >= 145, "Yes", "No")

    # Clean up and prepare final output
    aq_df = aq_df.drop_duplicates()
    aq_df = aq_df[['AQ_total', 'AQ_social', 'AQ_attention_switching', 'AQ_attention_to_detail', 'AQ_communication', 'AQ_imagination','AQ_above_diagnostic_threshold']]
    
    # Save to file
    output_file = project_folder / 'data'/ 'behavioral' / 'AQ_scores_Final.csv'
    aq_df.to_csv(output_file)
    print(f"AQ scores saved to: {output_file}")
    print(f"Processed {len(aq_df)} participants")
    
    return aq_df

# Process AQ questionnaire
aq_df = process_aq_questionnaire(data_folder, project_folder)

# %%
def process_sdq_questionnaire(data_folder, project_folder):
    """
    Process SDQ (Strengths and Difficulties Questionnaire) data.
    
    The SDQ is a behavioural screening questionnaire measuring emotional and
    behavioural problems in children and adolescents. It consists of 25 items
    across 5 subscales: emotional problems, conduct problems, hyperactivity,
    peer problems, and prosocial behaviour.
    
    Parameters:
    -----------
    data_folder : Path
        Path to the behavioural data folder
    project_folder : Path
        Path to the project root folder
        
    Returns:
    --------
    pd.DataFrame
        Processed SDQ scores with subscales and total score
    """
    print("Processing SDQ questionnaire data...")
    
    # Load and concatenate all SDQ files
    sdq_files = [f for f in data_folder.glob('**/SDQ*.csv')]
    print(f"Found {len(sdq_files)} SDQ files")
    
    sdq_df = pd.concat([pd.read_csv(f, skiprows=(1,2), index_col='ParticipantID', na_values='NaN') for f in sdq_files])
    sdq_df = sdq_df[[col for col in sdq_df.columns if 'SDQ' in col]].dropna().copy()
    sdq_df.columns = ['SDQ' + str(i+1) for i in range(len(sorted(sdq_df.columns)))]
    sdq_df = sdq_df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)
    
    # Replace ID that had a typo in original data
    sdq_df = sdq_df.rename(index={'TH1305225':'TH130525'})

    # Define forward and reverse scored items
    forward_items = [1, 2, 3, 4, 5, 6, 8, 9, 10, 12, 13, 15, 16, 17, 18, 19, 20, 22, 23, 24]
    reverse_items = [7, 11, 14, 21, 25]
    
    # Score forward items (higher problems = higher score)
    for item in forward_items:
        sdq_df.loc[:, 'SDQ' + str(item)] = sdq_df['SDQ' + str(item)].replace({
            'Niet waar':0,
            'Een beetje waar':1,
            'Zeker waar':2
            })
    
    # Score reverse items (positive behaviours - reverse scored)
    for item in reverse_items:
        sdq_df.loc[:, 'SDQ' + str(item)] = sdq_df['SDQ' + str(item)].replace({
            'Niet waar':2,
            'Een beetje waar':1,
            'Zeker waar':0
            })
    
    # Define subscale item groupings
    emotional_problems = [3, 8, 13, 16, 24]
    conduct_problems = [5, 7, 12, 18, 22]
    hyperactivity = [2, 10, 15, 21, 25]
    peer_problems = [6, 11, 14, 19, 23]
    prosocial = [1, 4, 9, 17, 20]
    
    # Calculate subscale and total scores
    sdq_df['SDQ_emotional_problems'] = sdq_df[[f'SDQ{i}' for i in emotional_problems]].sum(axis=1)
    sdq_df['SDQ_conduct_problems'] = sdq_df[[f'SDQ{i}' for i in conduct_problems]].sum(axis=1)
    sdq_df['SDQ_hyperactivity'] = sdq_df[[f'SDQ{i}' for i in hyperactivity]].sum(axis=1)
    sdq_df['SDQ_peer_problems'] = sdq_df[[f'SDQ{i}' for i in peer_problems]].sum(axis=1)
    sdq_df['SDQ_prosocial'] = sdq_df[[f'SDQ{i}' for i in prosocial]].sum(axis=1)
    sdq_df['SDQ_total'] = sdq_df[[f'SDQ{i}' for i in range(1, 26)]].sum(axis=1)
    
    # Clean up and prepare final output
    sdq_df = sdq_df.drop_duplicates()
    sdq_df = sdq_df[['SDQ_total', 'SDQ_emotional_problems', 'SDQ_conduct_problems', 'SDQ_hyperactivity', 'SDQ_peer_problems', 'SDQ_prosocial']]
    
    # Save to file
    output_file = project_folder / 'data'/ 'behavioral' / 'SDQ_scores_Final.csv'
    sdq_df.to_csv(output_file)
    print(f"SDQ scores saved to: {output_file}")
    print(f"Processed {len(sdq_df)} participants")
    
    return sdq_df

# Process SDQ questionnaire
sdq_df = process_sdq_questionnaire(data_folder, project_folder)

import os
import re
import pandas as pd

def process_arsq_questionnaire(data_folder, output_file):
    """
    Process ARSQ (Amsterdam Resting-State Questionnaire) data.
    
    Parameters
    ----------
    data_folder : str or Path
        Root directory containing participant subfolders with ARSQ CSV files.
        
    project_folder : Path  
        Path to the project root folder
    
    Returns
    -------
    pd.DataFrame
        - Long-format ARSQ dataframe with subscale dimension scores 
        for both version 1 (DOI: 10.3389/fnhum.2013.00446) and 
        version 2 (DOI: 10.3389/fpsyg.2014.00271) 
        - Wide-format ARSQ dataframe with subscale dimension scores
        for version 1 only
    """
    
    print("Processing ARSQ questionnaire data...")

    # 1. Dimension mappings
    dimension_V1 = {
        # Discontinuity of Mind
        "Ik had drukke gedachten.": "Discontinuity of Mind",
        "Ik had snel wisselende gedachten.": "Discontinuity of Mind",
        "Ik had moeite mijn gedachten vast te houden.": "Discontinuity of Mind",
        "Ik voelde me rusteloos.": "Discontinuity of Mind",
        "Ik had mijn gedachten onder controle.": "Discontinuity of Mind",

        # Theory of Mind
        "Ik dacht aan anderen.": "Theory of Mind",
        "Ik denk aan mensen die ik leuk vind.": "Theory of Mind",
        "Ik plaatste mezelf in de schoenen van anderen.": "Theory of Mind",

        # Self
        "Ik dacht aan mijn gevoelens.": "Self",
        "Ik dacht na over mijn gedrag.": "Self",
        "Ik dacht aan mezelf.": "Self",

        # Planning
        "Ik dacht aan mijn werk/studie.": "Planning",
        "Ik dacht na over de dingen die ik moet doen.": "Planning",
        "Ik dacht aan de toekomst.": "Planning",
        "Ik dacht aan het heden.": "Planning",
        "Ik dacht aan het oplossen van problemen.": "Planning",
        "Ik had diepe gedachten.": "Planning",
        "Ik dacht na over de dingen die ik moet doen.": "Planning",

        # Sleepiness
        "Ik voelde me moe.": "Sleepiness",
        "Ik voelde me slaperig.": "Sleepiness",
        "Ik had moeite om wakker te blijven.": "Sleepiness",

        # Comfort
        "Ik voelde me op mijn gemak.": "Comfort",
        "k voelde me ontspannen.": "Comfort",
        "Ik voelde me gelukkig.": "Comfort",

        # Somatic Awareness
        "Ik was me bewust van mijn lichaam.": "Somatic Awareness",
        "Ik dacht aan mijn hartslag.": "Somatic Awareness",
        "Ik dacht aan mijn ademhaling.": "Somatic Awareness",
        "Ik dacht aan mijn gezondheid.": "Somatic Awareness",

        # Validation
        "Ik had mijn ogen dicht.": "Validation",
        "Ik voelde mij gemotiveerd om mee te doen.": "Validation",
        "Ik heb moeite met het onthouden van mijn gedachten.": "Validation",
        "Ik heb moeite met het herinneren van mijn gevoelens.": "Validation",
        "Ik heb de uitspraken kunnen beoordelen.": "Validation",
    }

    dimension_V2 = {
        # Discontinuity of Mind
        "Ik had drukke gedachten.": "Discontinuity of Mind",
        "Ik had snel wisselende gedachten.": "Discontinuity of Mind",
        "Ik had moeite mijn gedachten vast te houden.": "Discontinuity of Mind",

        # Theory of Mind
        "Ik dacht aan anderen.": "Theory of Mind",
        "Ik denk aan mensen die ik leuk vind.": "Theory of Mind",
        "Ik plaatste mezelf in de schoenen van anderen.": "Theory of Mind",

        # Self
        "Ik dacht aan mijn gevoelens.": "Self",
        "Ik dacht na over mijn gedrag.": "Self",
        "Ik dacht aan mezelf.": "Self",

        # Planning
        "Ik dacht na over de dingen die ik moet doen.": "Planning",
        "Ik dacht aan het oplossen van problemen.": "Planning",
        "Ik dacht aan de toekomst.": "Planning",

        # Sleepiness
        "Ik voelde me moe.": "Sleepiness",
        "Ik voelde me slaperig.": "Sleepiness",
        "Ik had moeite om wakker te blijven.": "Sleepiness",

        # Comfort
        "Ik voelde me op mijn gemak.": "Comfort",
        "k voelde me ontspannen.": "Comfort",
        "Ik voelde me gelukkig.": "Comfort",

        # Somatic Awareness
        "Ik was me bewust van mijn lichaam.": "Somatic Awareness",
        "Ik dacht aan mijn hartslag.": "Somatic Awareness",
        "Ik dacht aan mijn ademhaling.": "Somatic Awareness",

        # Health Concern
        "Ik voelde me ziek.": "Health Concern",
        "Ik dacht aan mijn gezondheid.": "Health Concern",
        "Ik voelde pijn.": "Health Concern",

        # Visual Thought
        "Ik dacht in beelden.": "Visual Thought",
        "Ik stelde me gebeurtenissen voor.": "Visual Thought",
        "Ik stelde me plaatsen voor.": "Visual Thought",

        # Verbal Thought
        "Ik dacht in woorden.": "Verbal Thought",
        "Ik had stille gesprekken.": "Verbal Thought",
        "Ik stelde me voor dat ik met mezelf praatte.": "Verbal Thought",

        # Validation
        "Ik had mijn ogen dicht.": "Validation",
        "Ik voelde mij gemotiveerd om mee te doen.": "Validation",
        "Ik heb moeite met het onthouden van mijn gedachten.": "Validation",
        "Ik heb moeite met het herinneren van mijn gevoelens.": "Validation",
        "Ik heb de uitspraken kunnen beoordelen.": "Validation",
    }

    reverse_coding = [
        "Ik voelde pijn.",
        "Ik voelde me ziek.",
        "Ik had negatieve gevoelens.",
        "Ik had mijn gedachten onder controle."
    ]

    stimulus_to_condition = {
    "rest": "Rest",
    "FOMO": "Reality",
    "first_dates": "Reality",
    "interview" : "Reality",
    "polen": "News",
    "farmers" : "News",
    "royals": "News"
}

    # 2. Walk through files and build long dataframe
    long_data = []

    for root, dirs, files in os.walk(data_folder):
        for file in files:
            if (
                file.endswith(".csv")
                and (("VideoTask" in file) or ("RestingTask" in file))
                and not file.startswith("._")
            ):

                file_path = os.path.join(root, file)
                participant_id = os.path.basename(root).replace("sub-", "")

                # Extract session + stimulus
                match = re.search(r"_ses-(\d+)_stim-(\w+)\.csv$", file)
                if not match:
                    print(f"Skipping file with unmatched pattern: {file}")
                    continue

                ses, stim = match.group(1), match.group(2)

                # Load CSV
                df = pd.read_csv(file_path)

                if not set(["item", "Rating"]).issubset(df.columns):
                    print(f"Skipping {file}: missing columns")
                    continue

                df = df.dropna(subset=["item", "Rating"], how="all")

                # Dimension mapping
                df["Dimension_V1"] = df["item"].map(dimension_V1).fillna("Other")
                df["Dimension_V2"] = df["item"].map(dimension_V2).fillna("Other")

                # Reverse coding
                df["Rating_rc"] = df.apply(
                    lambda r: 5 - r["Rating"] if r["item"] in reverse_coding else r["Rating"],
                    axis=1
                )

                # Metadata
                df["Participant"] = participant_id
                df["Session"] = ses
                df["Stimulus"] = stim
                df["Condition"] = df["Stimulus"].map(stimulus_to_condition)

                long_data.append(df[[
                    "Participant", "Session", "Stimulus", "Condition",
                    "item", "Rating_rc",
                    "Dimension_V1", "Dimension_V2"
                ]])

    # 3. Combine and compute dimension sums
    arsq = pd.concat(long_data, ignore_index=True)
    arsq = arsq.rename(columns={"Rating_rc": "Rating", "item": "Item"})

    arsq["Sum_dimension_rating_V1"] = arsq.groupby(
        ["Participant", "Session", "Stimulus", "Dimension_V1"]
    )["Rating"].transform("sum")

    arsq["Sum_dimension_rating_V2"] = arsq.groupby(
        ["Participant", "Session", "Stimulus", "Dimension_V2"]
    )["Rating"].transform("sum")

    # 4. Save output in long format
    output_file = project_folder / 'data'/ 'behavioral' / 'ARSQ_scores_Dec25.csv'
    arsq.to_csv(output_file, index=False)
    print(f"ARSQ scores saved in long format to {output_file}")

    # 4. Save output in wide format
    arsq_v1_wide = (
        arsq.drop_duplicates(
            subset=["Participant", "Session", "Condition", "Dimension_V1"]
        )
        .pivot_table(
            index="Participant",
            columns=["Session", "Condition", "Dimension_V1"],
            values="Sum_dimension_rating_V1"
        )
    )

    arsq_v1_wide.columns = [
    f"Ses_{ses}_Condition_{cond}_Dimension_{dim.replace(' ', '_')}"
    for ses, cond, dim in arsq_v1_wide.columns
    ]
    
    output_file = project_folder / 'data'/ 'behavioral' / 'ARSQ_scores_Final_Wide.csv'
    arsq_v1_wide.to_csv(output_file)
    print(f"ARSQ scores saved in wide format to {output_file}")
    print(f"Processed {len(arsq_v1_wide)} participants")

    return arsq, arsq_v1_wide

# Process ARSQ questionnaire
arsq_long_df, arsq_df = process_arsq_questionnaire(data_folder, project_folder)

def process_wais_scores(data_folder, project_folder, norm_tables_excel, demographics="/home/ASDPrecision/participants.tsv"):
    """
    Process WAIS-IV scores:
    - Extract raw Matrix Reasoning and Vocabulary scores from .docx files
    - Write them out to .tsv files
    - Convert raw scores to scaled scores using norm tables
    - Save updated participants.tsv

    Parameters
    ----------
    data_folder : str
        Root folder containing raw behavioural data
    project_folder : Path  
        Path to the project root folder    
    norm_tables_excel : str
        Path to WAIS-IV Excel norm table
    demographics : str or Path
            Path to participants.tsv file
    
    Returns
    -------
    df : pd.DataFrame
        Updated WAIS IQ score table including: 
        - raw & scaled Matrix Reasoning
        - raw & scaled Vocabulary
        - PPVT indicator
    """
    
    # Utility functions
    def parse_age_range(sheet_name):
        low_str, high_str = sheet_name.split("-")
        low = float(low_str)
        high_year = high_str.split(".")[0]
        high = float(high_year) + 0.99
        return low, high

    def get_table_for_age(age, all_tables):
        for (low, high, table) in all_tables:
            if low <= age <= high:
                return table
        return None

    def raw_in_range(raw, cell_value):
        if pd.isna(cell_value) or cell_value == "":
            return False
        s = str(cell_value).strip()
        if re.fullmatch(r"\d+", s):
            return int(s) == raw
        if re.fullmatch(r"\d+-\d+", s):
            lo, hi = map(int, s.split("-"))
            return lo <= raw <= hi
        return False

    def get_scaled_score(table, raw_score, column_name):
        for _, row in table.iterrows():
            if raw_in_range(raw_score, row[column_name]):
                return int(row["Geschaalde score"])
        return np.nan       

    # Load demographics file
    iq = pd.read_csv(demographics, sep="\t")
    scores = {}

    # Extract tables from Word files
    for root, dirs, files in os.walk(data_folder):

        for file in files:

            # ---------------- Matrix Reasoning ----------------
            if (
                file.endswith(".docx")
                and "Matrix Reasoning Scoring Sheet" in file
                and "ERROR" not in file
                and not file.startswith("._")
            ):
                file_path = os.path.join(root, file)
                folder_name = os.path.basename(root)
                participant_id = folder_name.replace("sub-", "")

                doc = Document(file_path)
                table = doc.tables[0]

                stacked_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if not any(cells):
                        continue
                    if cells[0].lower() != "item number":
                        stacked_rows.append(cells[:3])
                    if len(cells) > 3 and cells[3].lower() != "item number":
                        stacked_rows.append(cells[3:6])

                def parse_item_number(x):
                    try:
                        return int(x.replace(".", "").strip())
                    except:
                        return float("inf")

                stacked_rows.sort(key=lambda r: parse_item_number(r[0]))

                # Score logic
                item_dict = {parse_item_number(r[0]): r for r in stacked_rows}
                has_later_scores = any(
                    r[2].isdigit() for num, r in item_dict.items() if num > 3
                )

                if has_later_scores:
                    matrix_score = sum(int(r[2]) for r in stacked_rows if r[2].isdigit())
                    for i in {1, 2, 3}:
                        if i in item_dict and not item_dict[i][2].isdigit():
                            matrix_score += 1
                else:
                    matrix_score = sum(int(r[2]) for r in stacked_rows if r[2].isdigit())

                scores[f"{participant_id}_matrix"] = matrix_score

            # ---------------- Vocabulary ----------------
            if (
                file.endswith(".docx")
                and "Woordenschat scoring sheet" in file
                and "ERROR" not in file
                and not file.startswith("._")
            ):

                file_path = os.path.join(root, file)
                folder_name = os.path.basename(root)
                participant_id = folder_name.replace("sub-", "")

                doc = Document(file_path)
                table = doc.tables[0]

                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                    if any(cell.text.strip() for cell in row.cells)
                ]

                header = rows[0]
                data_rows = rows[1:]

                vocab_score = sum(
                    int(r[2]) for r in data_rows if r[2].isdigit()
                )
                if vocab_score > 0:
                    vocab_score += 3

                scores[f"{participant_id}_vocab"] = vocab_score

    # Insert raw scores into iq dataframe
    iq["Vocabulary_raw_score"] = pd.NA
    iq["MatrixReasoning_raw_score"] = pd.NA
    
    for pid, score in scores.items():
        clean = pid.replace("_vocab", "").replace("_matrix", "")
        if pid.endswith("_vocab"):
            iq.loc[iq["ParticipantID"] == clean, "Vocabulary_raw_score"] = score
        if pid.endswith("_matrix"):
            iq.loc[iq["ParticipantID"] == clean, "MatrixReasoning_raw_score"] = score

    # PPVT participants
    PPVT_subs = ["AH070524","JH080524","FB250324","IB080524","AJ240524",
                 "MM190424","CV080424","NV110424","SB150424","EB310524","BP260324"]

    iq["CompletedPPVT"] = iq["ParticipantID"].apply(
        lambda x: "Yes" if x in PPVT_subs else "No"
    )

    # Load norm tables
    norm_tables = pd.ExcelFile(norm_tables_excel)
    tables = []

    for sheet_name in norm_tables.sheet_names:
        sheet = pd.read_excel(norm_tables_excel, sheet_name=sheet_name, header=1)
        sheet.columns = [c.replace("\n", " ").strip() for c in sheet.columns]
        low, high = parse_age_range(sheet_name)
        tables.append((low, high, sheet))

    # Convert raw to scaled scores
    iq["MatrixReasoning_scaled_score"] = np.nan
    iq["Vocabulary_scaled_score"] = np.nan

    for idx, row in iq.iterrows():
        age = row["Age"]
        mr_raw = row["MatrixReasoning_raw_score"]
        ws_raw = row["Vocabulary_raw_score"]

        table = get_table_for_age(age, tables)
        if table is None:
            print(f"No norm table for age {age} (sub-{row['ParticipantID']})")
            continue

        if pd.notna(mr_raw):
            iq.loc[idx, "MatrixReasoning_scaled_score"] = get_scaled_score(
                table, int(mr_raw), "MR"
            )

        if pd.notna(ws_raw):
            iq.loc[idx, "Vocabulary_scaled_score"] = get_scaled_score(
                table, int(ws_raw), "WS"
            )
    
    # Save output
    cols_to_save = [
        "ParticipantID",
        "MatrixReasoning_raw_score",
        "MatrixReasoning_scaled_score",
        "Vocabulary_raw_score",
        "Vocabulary_scaled_score",
        "CompletedPPVT"
    ]
    output_file = project_folder / 'data'/ 'behavioral' / 'WAIS_IQ_scores_Final.csv'
    iq = iq[cols_to_save].set_index("ParticipantID")
    iq.to_csv(output_file, index=False)
    print(f"WAIS IQ scores saved to: {output_file}")
    print(f"Processed {len(iq)} participants")

    return iq

# Process WAIS IQ scores 
iq_df = process_wais_scores(
    data_folder,
    project_folder,
    norm_tables_excel="/home/ASDPrecision/data/behavioral/WAIS_IV_Norm_conversion_tables.xlsx",
    demographics="/home/ASDPrecision/participants.tsv"
)
 # %%
def process_demographics(data_folder, project_folder):
    """
    Process demographics and participant information.
    
    This function processes demographic data including age calculation,
    handedness scoring, ethnicity extraction, and language/diagnosis information.
    
    Parameters:
    -----------
    data_folder : Path
        Path to the behavioural data folder
    project_folder : Path
        Path to the project root folder
        
    Returns:
    --------
    pd.DataFrame
        Processed demographics data with calculated variables
    """
    print("Processing demographics data...")
    
    # Load and concatenate all demographics files
    demographics_files = [f for f in data_folder.glob('**/Demographics*.csv')]
    print(f"Found {len(demographics_files)} demographics files")
    
    demographics_df = pd.concat([pd.read_csv(f, skiprows=(1,2), index_col='ParticipantID', na_values='NaN') for f in demographics_files])
    demographics_df = demographics_df.dropna(axis=1, how='all')

    # Replace index CV160425 with CV080425 (was wrongly labelled in original data)
    demographics_df = demographics_df.rename(index={'CV160424':'CV080424'})
    
    # Define relevant columns to keep
    relevant_columns = [
        'StartDate',
        'BirthDate#1_1',
        'BirthDate#2_1', 
        'Sex',
        'DutchFirstLang',
        'YearsEducation',
        'HandWriting',
        'HandDrawing',
        'HandPrecision',
        'Ethnicity',
        'AutismDiagnosis',
        'ASDDiagDate_1',
        'ADHDDiag',
        'ADHDDiagDate_1',
        'IntDisability',
        'OtherDiagnosis',
        'TypeDisorder',
        'OldOtherDisorder',
        'OldDisorderType',
        'FamilyASD',
        'FamilyADHD'
    ]
    
    # Filter participants with valid birth dates and select relevant columns
    demographics_df = demographics_df.loc[~demographics_df['BirthDate#1_1'].isna(), :].copy()
    demographics_df = demographics_df[relevant_columns].copy()
    
    # Calculate age from birth date and start date
    demographics_df['StartDate'] = pd.to_datetime(demographics_df['StartDate'], format='%Y-%m-%d %H:%M:%S')
    demographics_df['BirthDate'] = demographics_df['BirthDate#1_1'] + " " + demographics_df['BirthDate#2_1'].astype(int).astype(str)
    demographics_df['BirthDate'] = pd.to_datetime(demographics_df['BirthDate'], format='%B %Y')
    demographics_df['Age'] = (demographics_df['StartDate'] - demographics_df['BirthDate']).dt.days / 365.25
    demographics_df['Age'] = demographics_df['Age'].round(1)
    
    # Calculate handedness score (-1 = left-handed, 0 = ambidextrous, 1 = right-handed)
    handedness_mapping = {
        'Altijd rechts': 2,
        'Gewoonlijk rechts': 1,
        'Beide evenveel': 0,
        'Gewoonlijk links': -1,
        'Altijd links': -2
    }
    
    for handitem in ['HandWriting', 'HandDrawing', 'HandPrecision']:
        demographics_df.loc[:, handitem] = demographics_df[handitem].replace(handedness_mapping)
    
    demographics_df['Handedness'] = demographics_df[['HandWriting', 'HandDrawing', 'HandPrecision']].sum(axis=1)/6
    
    # Extract and simplify ethnicity information
    def extract_ethnicities_combined(entry):
        """Extract ethnicity information from complex ethnicity strings."""
        if pd.isna(entry):
            return 'Unknown'
        matches = ethnicity_regex.findall(entry)
        return ', '.join(matches) if matches else 'Unknown'
    
    ethnicities = ['Europa', 'Oost-Azië', 'Zuidoost-Azië', 'Zuid-Azië', 'West-Azië en Noord-Afrika', 
                   'West- en Centraal-Afrika', 'Oost- en Zuid-Afrika', 'Noord-Amerika', 'Zuid- en Midden-Amerika']
    ethnicity_pattern = '|'.join([re.escape(e) for e in ethnicities])
    ethnicity_regex = re.compile(rf'\b({ethnicity_pattern})\b', flags=re.IGNORECASE)
    demographics_df['MatchedEthnicities'] = demographics_df['Ethnicity'].apply(extract_ethnicities_combined)
    
    # Create a translation mapping for ethnicity terms
    ethnicity_translations = {
        'Europa': 'Europe',
        'Oost-Azië': 'East Asia',
        'Zuidoost-Azië': 'Southeast Asia',
        'Zuid-Azië': 'South Asia',
        'West-Azië en Noord-Afrika': 'West Asia and North Africa',
        'West- en Centraal-Afrika': 'West and Central Africa',
        'Oost- en Zuid-Afrika': 'East and South Africa',
        'Noord-Amerika': 'North America',
        'Zuid- en Midden-Amerika': 'South and Central America'
    }
    
    # Translate ethnicity entries from Dutch to English
    def translate_ethnicity(entry):
        """Translate Dutch ethnicity terms to English."""
        if pd.isna(entry):
            return entry
        translated_entry = entry
        for dutch_term, english_term in ethnicity_translations.items():
            translated_entry = translated_entry.replace(dutch_term, english_term)
        return translated_entry
    
    demographics_df['MatchedEthnicities'] = demographics_df['MatchedEthnicities'].apply(translate_ethnicity)

    # Translate Dutch responses to English
    translation_mappings = {
        'Sex': {'Mannelijk': 'Male', 'Vrouwelijk': 'Female'},
        'DutchFirstLang': {'Ja': 'Yes', 'Nee': 'No'},
        'YearsEducation': {
            '9-12 jaar (gelijk aan het voltooien van middelbare school, GCSE/A-level, VMBO/HAVO/VWO)': '9-12 years',
            '13-16 jaar (gelijk aan het voltooien van een bacheloropleiding)': '13-16 years',
            '17-20 jaar (gelijk aan het voltooien van een masteropleiding)': '17-20 years',
            'Meer dan 20 jaar (gelijk aan het voltooien van een doctoraat)': 'More than 20 years'
        },
        'AutismDiagnosis': {'Ja': 'Yes', 'Nee': 'No'},
        'ADHDDiag': {'Ja': 'Yes', 'Nee': 'No'},
        'IntDisability': {'Ja': 'Yes', 'Nee': 'No'},
        'OtherDiagnosis': {'Ja': 'Yes', 'Nee': 'No'}
    }
    
    for column, mapping in translation_mappings.items():
        demographics_df.loc[:, column] = demographics_df[column].replace(mapping)

    # Transform diagnosis date to only year
    demographics_df['ASDDiagDate_1'] = pd.to_datetime(demographics_df['ASDDiagDate_1'], errors='coerce', format='%Y')
    demographics_df['ASDDiagDate_1'] = demographics_df['ASDDiagDate_1'].dt.year
    demographics_df.rename(columns={'ASDDiagDate_1':'ASDDiagYear'}, inplace=True)

    # Define final output columns
    out_columns = [
        'Age',
        'Handedness',
        'Sex',
        'YearsEducation',
        'DutchFirstLang',
        'MatchedEthnicities',
        'AutismDiagnosis',
        'ASDDiagYear',
        'ADHDDiag',
        'IntDisability',
        'OtherDiagnosis'
    ]
    
    # Clean up and prepare final output
    demographics_df = demographics_df.drop_duplicates()
    demographics_df = demographics_df[out_columns]

    # Save to file
    output_file = project_folder / 'data'/ 'behavioral' / 'Demographics_Final.csv'
    demographics_df[out_columns].to_csv(output_file)
    print(f"Demographics data saved to: {output_file}")
    print(f"Processed {len(demographics_df)} participants")
    
    return demographics_df

# Process demographics
demographics_df = process_demographics(data_folder, project_folder)

# %%
def merge_all_data(demographics_df, aq_df, sdq_df, iq_df, arsq_df, project_folder, acquisition_dates = "/home/ASDPrecision/data/behavioral/acquisition_dates.csv"):
    """
    Merge all processed datasets into a single comprehensive dataset.
    
    This function combines demographics and questionnaire scores into a 
    single dataset for analysis.
    
    Parameters:
    -----------
    demographics_df : pd.DataFrame
        Processed demographics data
    aq_df : pd.DataFrame
        Processed AQ scores
    sdq_df : pd.DataFrame
        Processed SDQ scores
    iq_df : pd.DataFrame
        Processed WAIS IQ scores for Matrix Reasoning and Vocabulary
    arsq_df : pd.DataFrame
        Processed ARSQ scores in wide format
    project_folder : Path
        Path to the project root folder
    acquisition_dates : Path
        Dataframe of dates of MRI scans
        
    Returns:
    --------
    pd.DataFrame
        Merged dataset with all participant information
    """
    print("Merging all datasets...")
    
    # Start with demographics as the base
    merged_df = demographics_df.copy()
    print(f"Starting with {len(merged_df)} participants from demographics")

    # Merge AQ scores
    merged_df = pd.merge(merged_df, aq_df, left_index=True, right_index=True, how='left').drop_duplicates()
    print(f"After merging AQ: {len(merged_df)} participants, {aq_df.notna().any(axis=1).sum()} with AQ data")
    
    # Merge SDQ scores
    merged_df = pd.merge(merged_df, sdq_df, left_index=True, right_index=True, how='left').drop_duplicates()
    print(f"After merging SDQ: {len(merged_df)} participants, {sdq_df.notna().any(axis=1).sum()} with SDQ data")

    # Merge WAIS IQ scores
    merged_df = pd.merge(merged_df, iq_df, left_index=True, right_index=True, how='left').drop_duplicates()
    print(f"After merging WAIS IQ: {len(merged_df)} participants, {iq_df.notna().any(axis=1).sum()} with WAIS IQ data")

    # Merge ARSQ scores
    merged_df = pd.merge(merged_df, arsq_df, left_index=True, right_index=True, how='left').drop_duplicates()
    print(f"After merging ARSQ: {len(merged_df)} participants, {arsq_df.notna().any(axis=1).sum()} with ARSQ data")

    # Merge acquisition dates
    acquisition_date_file = project_folder / 'data'/ 'behavioral' / 'acquisition_dates.csv'
    acquisition_dates = pd.read_csv(acquisition_date_file).set_index("subject")
    merged_df = pd.merge(merged_df, acquisition_dates, left_index=True, right_index=True, how='left').drop_duplicates()
    print(f"After merging ARSQ: {len(merged_df)} participants, {acquisition_dates.notna().any(axis=1).sum()} with acquisition date data")

    # Final cleanup
    merged_df['Excluded'] = 0
    merged_df.loc['EB310524', 'Excluded'] = 1 # Excluded because of age (47)
    merged_df.loc['OD051124', 'Excluded'] = 1 # Excluded because of piercings
    merged_df.loc['PW200225', 'Excluded'] = 1 # Excluded because of scanner issues
    merged_df.loc['RO140325', 'Excluded'] = 1 # Withdrew from study
    merged_df.loc['LV161024', 'Excluded'] = 1 # Never completed an MRI session
    merged_df = merged_df.sort_index()
    merged_df = merged_df[merged_df['Excluded'] == 0]
    merged_df = merged_df.drop(columns=['Excluded'])
    print(f"After removing excluded participants: {len(merged_df)} participants remain")


    # Renaming participants
    merged_df.index = 'sub-' + merged_df.index.astype(str)
    participant_renaming_file = project_folder / 'data'/ 'participant_renaming.tsv'
    participant_renaming = pd.read_csv(participant_renaming_file, sep = "\t").set_index("participant_id")
    merged_df = merged_df = pd.merge(merged_df, participant_renaming, left_index=True, right_index=True, how='left').drop_duplicates()
    merged_df = merged_df.reset_index(drop=True)  
    merged_df = merged_df.set_index('new_participant_id')
    merged_df = merged_df.rename_axis('ParticipantID')

    # Save merged dataset
    output_file = project_folder / 'data'/ 'behavioral' / 'Merged_Behavioural_Data_Final.csv'
    merged_df.to_csv(output_file)
    print(f"Merged dataset saved to: {output_file}")
    print(f"Final dataset contains {len(merged_df)} participants with {merged_df.shape[1]} variables")
    
    # Print summary statistics
    print("\n=== Data Summary ===")
    print(f"Total participants: {len(merged_df)}")
    print(f"Participants with complete demographics: {demographics_df.drop_duplicates().notna().all(axis=1).sum()}")
    print(f"Participants with AQ data: {merged_df[['AQ_total']].notna().sum().iloc[0] if 'AQ_total' in merged_df.columns else 0}")
    print(f"Participants with SDQ data: {merged_df[['SDQ_total']].notna().sum().iloc[0] if 'SDQ_total' in merged_df.columns else 0}")
    print(f"Participants with WAIS Matrix Reasoning data: {merged_df[['MatrixReasoning_raw_score']].notna().sum().iloc[0] if 'MatrixReasoning_raw_score' in merged_df.columns else 0}")
    print(f"Participants with WAIS Vocabulary data: {merged_df[['Vocabulary_raw_score']].notna().sum().iloc[0] if 'Vocabulary_raw_score' in merged_df.columns else 0}")

    return merged_df

# Merge all datasets
final_dataset = merge_all_data(demographics_df, aq_df, sdq_df, iq_df, arsq_df, project_folder)

print("\n=== Processing Complete ===")
print("All behavioural data has been processed and merged successfully!")
print(f"Final dataset shape: {final_dataset.shape}")
print(f"Output files saved in: {project_folder}")

# %%
# Comparison of MRI data and the final merged dataset
print("Comparing MRI data with final merged dataset...")
for ses in [1, 2, 3]:
    ses_folder = project_folder / 'data' / 'raw_data' / f'ses-0{ses}'
    ses_subjects = sorted([sub.replace('sub-', '') for sub in os.listdir(ses_folder) if sub.startswith('sub-')])
    print(f"Total MRI subjects in ses-0{ses}: {len(ses_subjects)}")
    mri_only = set(ses_subjects) - set(final_dataset.index)
    final_only = set(final_dataset.index) - set(ses_subjects)
    print(f"Subjects in MRI data but not in final dataset (ses-0{ses}): {mri_only}")
    print(f"Subjects in final dataset but not in MRI data (ses-0{ses}): {final_only}")

# %%

